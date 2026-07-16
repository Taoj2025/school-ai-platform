"""
Generate DOCX templates for Apple subsystem.
Run this script to create the certificate templates.

Requirements:
    pip install docxtpl

Usage:
    python generate_templates.py
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

def create_certificate_template():
    """Create the award certificate template."""
    doc = Document()
    
    # Set page margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.5)
        section.right_margin = Inches(1.5)
    
    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('獎  狀')
    run.bold = True
    run.font.size = Pt(36)
    run.font.name = '標楷體'
    
    # School name
    school = doc.add_paragraph()
    school.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = school.add_run('香港培英中學')
    run.font.size = Pt(20)
    run.font.name = '標楷體'
    
    doc.add_paragraph()
    
    # Main content
    content = doc.add_paragraph()
    content.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = content.add_run('茲證明')
    run.font.size = Pt(18)
    run.font.name = '標楷體'
    
    # Student name placeholder
    name = doc.add_paragraph()
    name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = name.add_run('{{ student_name }}')
    run.bold = True
    run.font.size = Pt(28)
    run.font.name = '標楷體'
    
    # Class placeholder
    cls = doc.add_paragraph()
    cls.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cls.add_run('（學號：{{ student_no }}，{{ class }}班）')
    run.font.size = Pt(14)
    run.font.name = '標楷體'
    
    doc.add_paragraph()
    
    # Award content
    award_content = doc.add_paragraph()
    award_content.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = award_content.add_run('品學兼優，獲頒')
    run.font.size = Pt(18)
    run.font.name = '標楷體'
    
    # Award name placeholder
    award_name = doc.add_paragraph()
    award_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = award_name.add_run('「{{ award_name }}」')
    run.bold = True
    run.font.size = Pt(24)
    run.font.name = '標楷體'
    
    doc.add_paragraph()
    
    # Date
    date = doc.add_paragraph()
    date.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = date.add_run('中華民國 {{ year }}年 {{ month }}月 {{ day }}日')
    run.font.size = Pt(14)
    run.font.name = '標楷體'
    
    # Principal
    principal = doc.add_paragraph()
    principal.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = principal.add_run('香港培英中學校長　{{ principal_name }}')
    run.font.size = Pt(14)
    run.font.name = '標楷體'
    
    doc.save('apps/api/templates/apple/certificate.docx')
    print('Created: apps/api/templates/apple/certificate.docx')


def create_student_certificate_template():
    """Create the student enrollment certificate template."""
    doc = Document()
    
    # Set page margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.5)
        section.right_margin = Inches(1.5)
    
    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('在 學 證 明 書')
    run.bold = True
    run.font.size = Pt(32)
    run.font.name = '標楷體'
    
    doc.add_paragraph()
    
    # School info
    school = doc.add_paragraph()
    school.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = school.add_run('香港培英中學')
    run.font.size = Pt(18)
    run.font.name = '標楷體'
    run.underline = True
    
    doc.add_paragraph()
    
    # Main content in Chinese
    zh_content = doc.add_paragraph()
    zh_content.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = zh_content.add_run(
        '茲證明下列學生為本校在讀學生：\n\n'
        '學生姓名：{{ student_name }}\n'
        '學　　號：{{ student_no }}\n'
        '班　　別：{{ class }}\n'
        '入學日期：{{ enrollment_date }}'
    )
    run.font.size = Pt(16)
    run.font.name = '標楷體'
    
    doc.add_paragraph()
    
    # Purpose
    purpose = doc.add_paragraph()
    purpose.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = purpose.add_run('本證明書僅作{{ purpose }}之用。')
    run.font.size = Pt(14)
    run.font.name = '標楷體'
    
    doc.add_paragraph()
    
    # Issue date
    issue_date = doc.add_paragraph()
    issue_date.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = issue_date.add_run('發出日期：{{ issue_date }}')
    run.font.size = Pt(14)
    run.font.name = '標楷體'
    
    doc.add_paragraph()
    
    # Principal
    principal = doc.add_paragraph()
    principal.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = principal.add_run('香港培英中學校長')
    run.font.size = Pt(14)
    run.font.name = '標楷體'
    
    doc.add_paragraph()
    
    # Principal name
    principal_name = doc.add_paragraph()
    principal_name.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = principal_name.add_run('（{{ principal_name }}）')
    run.font.size = Pt(14)
    run.font.name = '標楷體'
    
    # Separator line
    doc.add_paragraph()
    doc.add_paragraph('─' * 40)
    
    # English version
    doc.add_paragraph()
    
    en_title = doc.add_paragraph()
    en_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = en_title.add_run('CERTIFICATE OF ENROLMENT')
    run.bold = True
    run.font.size = Pt(24)
    run.font.name = 'Arial'
    
    en_school = doc.add_paragraph()
    en_school.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = en_school.add_run('Hong Kong Pui Ying Secondary School')
    run.font.size = Pt(16)
    run.font.name = 'Arial'
    run.underline = True
    
    doc.add_paragraph()
    
    # English content
    en_content = doc.add_paragraph()
    en_content.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = en_content.add_run(
        'This is to certify that {{ student_name_en }},\n'
        'Student Number: {{ student_no }},\n'
        'Class: {{ class }},\n'
        'has been enrolled in our school since {{ enrollment_date }}.'
    )
    run.font.size = Pt(14)
    run.font.name = 'Arial'
    
    doc.add_paragraph()
    
    # English purpose
    en_purpose = doc.add_paragraph()
    en_purpose.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = en_purpose.add_run('This certificate is issued for the purpose of {{ purpose_en }}.')
    run.font.size = Pt(12)
    run.font.name = 'Arial'
    
    doc.add_paragraph()
    
    # English date
    en_date = doc.add_paragraph()
    en_date.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = en_date.add_run('Date of Issue: {{ issue_date }}')
    run.font.size = Pt(12)
    run.font.name = 'Arial'
    
    doc.add_paragraph()
    
    # English principal
    en_principal = doc.add_paragraph()
    en_principal.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = en_principal.add_run('Principal: {{ principal_name_en }}')
    run.font.size = Pt(12)
    run.font.name = 'Arial'
    
    doc.save('apps/api/templates/apple/student_certificate.docx')
    print('Created: apps/api/templates/apple/student_certificate.docx')


if __name__ == '__main__':
    print('Generating Apple subsystem templates...')
    print()
    
    try:
        create_certificate_template()
        create_student_certificate_template()
        print()
        print('All templates generated successfully!')
    except ImportError:
        print('Error: docxtpl library not found.')
        print('Please install it with: pip install docxtpl')
        print()
        print('Alternative: Copy the template structure below to create the files manually.')
        print()
        print('certificate.docx should contain:')
        print('  - 獎 狀 (title)')
        print('  - 香港培英中學')
        print('  - 茲證明 [student_name]')
        print('  - 學號：[student_no]，班：[class]')
        print('  - 品學兼優，獲頒「[award_name]」')
        print('  - 日期：[year]年[month]月[day]日')
        print('  - 香港培英中學校長 [principal_name]')
        print()
        print('student_certificate.docx should contain:')
        print('  - 在學證明書')
        print('  - 香港培英中學')
        print('  - 學生信息：中英文')
        print('  - 發出日期')
        print('  - 校長簽名')
